"""Brazil NIPPT Word document extraction. Adapted from extra_info_air_20260323.py."""
import re
import os
import datetime


def extract_brazil_docx(file_path, source="巴西万基"):
    def _safe_get(lst, idx, default=""):
        """Safely get list element, returning default if out of bounds."""
        try:
            return lst[idx]
        except (IndexError, TypeError):
            return default

    """Extract data from a Brazil NIPPT registration Word document.
    Returns a dict of Sample fields or None on failure."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for Brazil document extraction")

    doc = docx.Document(file_path)

    # ── Table 0: seq + client_code ──
    table0_cells = []
    for row in doc.tables[0].rows:
        for cell in row.cells:
            table0_cells.append(cell.text.strip())

    seq, client_code = _extract_seq_code(table0_cells)
    seq = seq.replace('Cliente nº ', '').replace(" ", "").replace("_", "").replace('Clientenº:', '')

    # Extract source code from client_code (e.g. "VGBRDBM260529215" -> "DBM")
    source_match = re.search(r'^VGBR([A-Z]+)\d+$', client_code)
    src_code = source_match.group(1) if source_match else ""
    sales, agent = _assign_sales_agent(seq, src_code)

    # Tube IDs
    tubo_match = re.search(r"(\d+)\d*$", client_code)
    tubo = tubo_match.group(1) if tubo_match else ""

    # ── Tables 1 & 3: test info, collection location, etc. ──
    table13 = []
    for t in doc.tables:
        if t.rows[0].cells[0].text in ("Preço do teste:", "Local da coleta:"):
            for row in t.rows:
                for cell in row.cells:
                    table13.append(cell.text.strip() if cell.text.strip() else "")

    # Price, Sinal, Balance
    price = _safe_get(table13, 1).replace("R$", "").strip() if len(table13) > 1 else ""
    sinal = _safe_get(table13, 3).replace("R$", "").strip() if len(table13) > 3 else ""
    balance = _safe_get(table13, 5).replace("R$", "").strip() if len(table13) > 5 else ""

    # Report due date
    if len(table13) > 15:
        report_due_raw = "".join(re.findall(r"resultados(.+?)$", _safe_get(table13, 15)))
        report_due_date = report_due_raw.replace("(DD/MM/YY)", "").replace(":", "").strip()
    else:
        report_due_date = 

    # Client name
    if len(table13) > 9:
        client_name = _safe_get(table13, 9).replace(" De ", " de ").replace(" Dos ", " dos ").replace(" Da ", " da ")
    else:
        client_name = ""

    # Working days
    wd = _extract_days(_safe_get(table13, 15)) if len(table13) > 15 else None

    # Collection location
    collect_locate = _safe_get(table13, -7) if len(table13) >= 7 else ""
    collect_locate = _replace_collect_locate(collect_locate)

    # Gestational week
    gestational_week_raw = _safe_get(table13, 38) if len(table13) > 38 else ""

    # ── Table with "Nome Gestante": patient info ──
    table2 = []
    for t in doc.tables:
        if t.rows[0].cells[0].text == "Nome Gestante":
            for ri, row in enumerate(t.rows):
                if ri % 2 != 0:  # odd rows have the data
                    for cell in row.cells:
                        table2.append(cell.text.strip() if cell.text.strip() else "")

    fe_name = _safe_get(table2, 0).replace(" De ", " de ").replace(" Dos ", " dos ").replace(" Da ", " da ") if table2 else ""
    fe_name = fe_name.replace("+", " / ").replace(" fta", " FTA")
    rg_f = _safe_get(table2, 1)
    date_of_birth = _safe_get(table2, 2)

    # Sample type
    f_type_idx = 9 if len(table2) == 12 else -3
    f_type = table2[f_type_idx] if table2 else ""
    f_type = f_type.replace("periferico", "").replace(" ", "").replace("SANGUE", "Sangue").replace("SWAB", "Swab").replace("/", " / ").replace("+", " / ").replace("fta", "FTA")

    # Collection date (孕妇采样日期)
    col_date_idx = 4 if len(table2) == 12 else 3
    collection_date_raw = table2[col_date_idx].replace("/24", "/2024") if len(table2) > col_date_idx else ""

    # Doubting father name
    df_idx = 6 if len(table2) == 12 else 5
    doubting_father = table2[df_idx].title() if len(table2) > df_idx else ""
    doubting_father = doubting_father.replace(" De ", " de ").replace(" Dos ", " dos ").replace(" Da ", " da ").replace("+", " / ")
    if re.match(r'[A-Za-z]+\d+', doubting_father):
        doubting_father = doubting_father.upper()

    # ── Fetal gender ──
    fetal_gender = _extract_fetal_gender(doc)

    # ── Test item ──
    # Check seq suffix first, then fall back to docx content hints
    if seq.endswith("NP"):
        test_option = "Plus"
    elif seq.endswith("NB"):
        test_option = "Basic"
    else:
        # Scan all table text for NIPT type indicators
        all_text = " ".join(cell.text.upper() for t in doc.tables for row in t.rows for cell in row.cells)
        if "EXPANDIDO" in all_text or "EXPANDED" in all_text:
            test_option = "Plus"
        else:
            test_option = "Basic"

    # ── Map to Sample fields ──
    info = {
        'sample_source': '巴西万基',
        'test_option': test_option,
        'external_id': client_code,
        'collection_date_raw': collection_date_raw,
        'collection_date': _parse_date_br(collection_date_raw),
        'ordering_facility': collect_locate,
        'ordering_physician': '',
        'id_card': rg_f,
        'patient_name': client_name if client_name else fe_name,
        'patient_dob_raw': date_of_birth,
        'patient_dob': _parse_date_br(date_of_birth),
        'gestational_weeks_raw': gestational_week_raw,
        'gestational_weeks': _parse_int(gestational_week_raw) if gestational_week_raw else None,
        'age': None,
        'multiple_gestation': False,
        'ivf_status': False,
        'clinical_diagnosis': '',
        'price': price,
        'sinal': sinal,
        'balance': balance,
        'fetal_gender': fetal_gender,  # Brazil-specific
        'seq': seq,
        'sales': sales,
        'agent': agent,
        'tubo_f': tubo + "-F",
        'tubo_m': tubo + "-M",
        'doubting_father': doubting_father,
        'sample_type_raw': f_type,
        'report_due_date': report_due_date,
        'wd': f"{wd} WD" if wd else "",
    }
    return info


def _extract_seq_code(table0):
    seq, client_code = "", ""
    if len(table0) > 1 and table0[1].strip():
        seq = table0[1].strip()
    elif table0:
        seq = table0[0].replace("COLETA:", "").replace("CLIENTE:", "").strip()

    if len(table0) > 3 and table0[3].strip():
        client_code = table0[3].strip()
    elif len(table0) > 2 and table0[2].strip():
        client_code = table0[2].replace("Código do exame:", "").replace(".", "").strip()

    return seq, client_code


def _extract_days(text):
    match = re.search(r"[_\s]*(\d+)[_\s]*dias\s+úteis", text, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _assign_sales_agent(seq, source):
    sales, agent = "", ""
    if seq and str(seq).strip():
        if str(seq).isdigit():
            sales = source
        else:
            agent = source
    return sales, agent


def _replace_collect_locate(collect_locate):
    s = collect_locate.lower()
    has_rua = "rua" in s
    has_37cj11 = "37-cj11" in s
    if has_rua and not has_37cj11:
        return "Home Collection"
    elif has_rua and has_37cj11:
        return "VANGENES, São Paulo, SP"
    return collect_locate


def _extract_fetal_gender(doc):
    """Extract fetal gender matching original extra_info_air script."""
    fetal_info = ""
    if len(doc.tables) == 4:
        temp = []
        for row in doc.tables[1].rows:
            for cell in row.cells:
                temp.append(cell.text.strip())
        if "Sexagem Fetal" in str(temp) or "Sexagem:" in str(temp):
            fetal_info = temp[32] if len(temp) > 32 else ""
    elif len(doc.tables) == 5:
        temp = []
        for row in doc.tables[2].rows:
            for cell in row.cells:
                temp.append(cell.text.strip())
        if "Sexagem Fetal" in str(temp) or "Sexagem:" in str(temp):
            fetal_info = temp[1] if len(temp) > 1 else ""
    else:
        # Fallback for other structures
        table_idx = 1 if len(doc.tables) >= 2 else None
        if table_idx is not None:
            temp = []
            for row in doc.tables[table_idx].rows:
                for cell in row.cells:
                    temp.append(cell.text.strip())
            if "Sexagem Fetal" in str(temp) or "Sexagem:" in str(temp):
                fetal_info = temp[32] if len(temp) > 32 else (temp[1] if len(temp) > 1 else "")

    # Parse the gender info
    m1 = re.search(r"(\(\s*[xX]\s*\)\s*[Ss][Ii][Mm]\s*)\s*/", fetal_info)
    m1_2 = re.search(r"(\s*[Ss][Ii][Mm]\s*\(\s*[xX]\s*\))\s*/", fetal_info)
    m2 = re.search(r"/\s*(\(\s*[xX]\s*\)\s*[Nn][ãÃ][Oo])", fetal_info)
    m2_2 = re.search(r"/\s*[Nn][ãÃ][Oo]\s*(\(\s*[xX]\s*\)\s*)", fetal_info)

    if m1 or m1_2:
        return "Yes"
    elif m2 or m2_2:
        return "No"
    return "No"


def _parse_date_br(date_str):
    """Parse DD/MM/YYYY to YYYY-MM-DD."""
    if not date_str:
        return None
    for fmt in ["%d/%m/%Y", "%d/%m/%y"]:
        try:
            return datetime.datetime.strptime(date_str.strip(), fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return None


def _parse_int(val):
    """Safely parse an integer from strings like '12 SEMANAS'."""
    if not val:
        return None
    import re
    m = re.search(r'\d+', str(val))
    if m:
        return int(m.group())
    try:
        return int(str(val).strip())
    except (ValueError, TypeError):
        return None


def extract_docxs_from_directory(directory, source="巴西万基"):
    """Extract all .docx files in a directory."""
    results = []
    for root, dirs, files in os.walk(directory):
        for fname in sorted(files):
            if fname.lower().endswith('.docx') and not fname.startswith('.') and '~$' not in fname:
                fpath = os.path.join(root, fname)
                print(f"Processing: {fname}")
                try:
                    info = extract_brazil_docx(fpath, source)
                    if info:
                        results.append(info)
                        print(f"  -> {info.get('patient_name', 'N/A')} | {info.get('test_option', 'N/A')}")
                except Exception as e:
                    print(f"  -> FAILED: {e}")
    return results


def generate_excel_brazil(samples, output_path):
    """Generate Excel matching 调整结果.xlsx format (38 columns with header row)."""
    import pandas as pd
    import numpy as np

    rows = []
    for s in samples:
        row = {
            'Sample_source': '巴西万基',
            'Test_Option': s.get('test_option', ''),
            'Accessioning_ID': s.get('external_id', ''),
            'Collection_Date': s.get('collection_date_raw', ''),
            'Acceptance_Date': '',  # Manual
            'Hospital_or_Clinic': s.get('ordering_facility', ''),
            'Physician': '',
            'Patient_ID': s.get('id_card', ''),
            'Name': s.get('patient_name', ''),
            'DOB': s.get('patient_dob_raw', ''),
            'Gestational_Week': s.get('gestational_weeks', ''),
            'Report_code': '',
            'send_report_id': s.get('external_id', ''),
            'Age': s.get('age', np.nan),
            'Last_Menstrual_Period': np.nan,
            'Single/Twin': '-',
            'IVF': np.nan,
            'Pregnancy_History': np.nan,
            'Previous_Medical_History': np.nan,
            'FedEx_No.': np.nan,
            'Zscore21': np.nan, 'Zscore18': np.nan, 'Zscore13': np.nan,
            'T21': np.nan, 'T18': np.nan, 'T13': np.nan,
            'XO': np.nan, 'XXX': np.nan, 'XXY': np.nan, 'XYY': np.nan,
            'All_chrom': np.nan, 'ff': np.nan,
            'Gender_report': s.get('fetal_gender', np.nan),
            'Gender': np.nan,
            'Actua_Report_Date': np.nan, 'Report_Date': np.nan,
            'Unnamed: 36': np.nan, 'Report_situation': np.nan,
        }
        rows.append(row)

    df = pd.DataFrame(rows)
    df = df.sort_values(by=['Test_Option', 'Accessioning_ID'])
    df.to_excel(output_path, index=False)
    print(f"Excel saved to: {output_path}")
    return output_path
