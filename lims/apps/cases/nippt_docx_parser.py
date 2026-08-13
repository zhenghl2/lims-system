# -*- coding: utf-8 -*-
"""
NIPPT 巴西送检单（FICHA CLIENTE docx）解析模块
移植自 extra_info_air_20260813.py 的提取逻辑，输出与 batch_import_nippt 接口一致的 case 数据。

用法:
    from nippt_docx_parser import parse_docx_bytes
    result = parse_docx_bytes(docx_bytes)   # -> dict 或抛 NipptDocxParseError
"""
import io
import re

import docx

SAMPLE_TYPE_MAP = {
    "SANGUE": "BLOOD",
    "FTA": "DBS",
    "SWAB": "SWAB",
    "CABELO": "HAIR",
    "UNHA": "NAIL",
    "SEMEN": "SEMEN",
    "ESCOVA": "TOOTHBRUSH",
}


class NipptDocxParseError(Exception):
    pass


def extract_seq_code(table0):
    """从 table0 中提取 seq 和 cilent_code"""
    seq, cilent_code = "", ""
    if len(table0) > 1 and table0[1].strip():
        seq = table0[1].strip()
    else:
        seq = table0[0].replace("COLETA:", "").replace("CLIENTE:", "").strip()

    if len(table0) > 3 and table0[3].strip():
        cilent_code = table0[3].strip()
    elif len(table0) > 2 and table0[2].strip():
        cilent_code = table0[2].replace("Código do exame:", "").replace(".", "").strip()

    return seq, cilent_code


def extract_days(text):
    """提取 'xx dias úteis' 中的天数"""
    if not text:
        return None
    match = re.search(r"[_\s]*(\d+)[_\s]*dias\s+úteis", text, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def replace_collect_locate(collect_locate):
    s = (collect_locate or "").lower()
    has_rua = "rua" in s
    has_37cj11 = "37-cj11" in s
    if has_rua and not has_37cj11:
        return "Home Colletion"
    elif has_rua and has_37cj11:
        return "VANGENES, São Paulo, SP"
    return collect_locate or ""


def get_test_item(observacoes, seq):
    """NIPPT/NIPT 判定：Observações 里出现 NIPT XXX 按 XXX 定档；否则 seq 以 NP/NB 结尾"""
    if observacoes:
        m = re.search(r"NIPT\s*[-–|/:]?\s*(B[ÁÀA]SICO|PLUS)", observacoes, re.IGNORECASE)
        if m:
            return "Plus" if m.group(1).upper() == "PLUS" else "Basic"
    if seq and seq.endswith("NP"):
        return "Plus"
    if seq and seq.endswith("NB"):
        return "Basic"
    return "nippt"


def title_pt(name):
    return (name or "").title().replace(" De ", " de ").replace(" Dos ", " dos ").replace(" Da ", " da ")


def format_sample_type(s):
    s = (s or "").replace("periferico", "").replace(" ", "")
    s = s.replace("SANGUE", "Sangue").replace("SWAB", "Swab").replace("/", " / ").replace("+", " / ")
    s = s.replace("fta", "FTA")
    return s


def map_sample_types(s):
    """'Sangue / FTA' -> ['BLOOD', 'DBS']"""
    s = format_sample_type(s)
    if not s:
        return ["BLOOD"]
    result = []
    for part in s.split("/"):
        part = part.strip().upper()
        if part in SAMPLE_TYPE_MAP:
            result.append(SAMPLE_TYPE_MAP[part])
    return result if result else ["BLOOD"]


def extract_gestational_week(table13):
    """按 'Idade Gestacional' 标签定位孕周值；返回周数 int 或 None"""
    for i, cell in enumerate(table13):
        if (cell or "").strip().lower() == "idade gestacional":
            for j in range(i + 1, min(i + 10, len(table13))):
                v = table13[j].strip() if table13[j] else ""
                if v and v.lower() != "idade gestacional":
                    m = re.search(r"(\d+)\s*semana", v, re.IGNORECASE)
                    if m:
                        return int(m.group(1))
                    m2 = re.match(r"(\d+)", v)
                    if m2:
                        return int(m2.group(1))
            return None
    return None


def safe_get(cells, idx):
    return cells[idx] if len(cells) > idx else ""


def parse_people_table(table):
    """按表头行解析人员表：返回 (孕妇cells, [疑父cells,...])"""
    mother = None
    fathers = []
    mode = None
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        first = cells[0] if cells else ""
        if first.startswith("Nome Gestante"):
            mode = "mother"
            continue
        if first.startswith("Nome do Suposto Pai"):
            mode = "father"
            continue
        if mode == "mother" and first and mother is None:
            mother = cells
        elif mode == "father" and first:
            fathers.append(cells)
    if mother is None:
        mother = [""] * 6
    return mother, fathers


def extract_gender_info(doc):
    """性别鉴定提取（4/5 表格两种变体）"""
    fetal_gender_info = ""
    if len(doc.tables) == 4:
        temp = []
        for row in doc.tables[1].rows:
            for cell in row.cells:
                temp.append(cell.text.strip())
        if "Sexagem Fetal" in temp or "Sexagem:" in temp:
            fetal_gender_info = temp[32] if len(temp) > 32 else ""
    if len(doc.tables) == 5:
        temp = []
        for row in doc.tables[2].rows:
            for cell in row.cells:
                temp.append(cell.text.strip())
        if "Sexagem Fetal" in temp or "Sexagem:" in temp:
            fetal_gender_info = temp[1] if len(temp) > 1 else ""

    match1 = re.search(r"(\(\s*[xX]\s*\)\s*[Ss][Ii][Mm]\s*)\s*/", fetal_gender_info)
    match1_2 = re.search(r"(\s*[Ss][Ii][Mm]\s*\(\s*[xX]\s*\))\s*/", fetal_gender_info)
    match2 = re.search(r"/\s*(\(\s*[xX]\s*\)\s*[Nn][ãÃ][Oo])", fetal_gender_info)
    match2_2 = re.search(r"/\s*[Nn][ãÃ][Oo]\s*(\(\s*[xX]\s*\)\s*)", fetal_gender_info)

    if match1 or match1_2:
        return "Yes"
    return "No"


def parse_docx_bytes(data, filename=""):
    """解析一个 docx 文件字节流，返回分组后的 case dict（schema 同 batch_import_nippt）。

    返回 dict:
        {"seq", "mother_name", "mother_dob", "mother_id_card",
         "gestational_age_weeks", "collection_date", "fathers",
         "sales_person", "price", "balance", "gender_info", "notes",
         "expected_completion"}
    或 {"_nipt": True, "test_item": "Basic|Plus"}（NIPT 送检单）
    解析失败抛 NipptDocxParseError。
    """
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise NipptDocxParseError(f"{filename}: 无法打开 Word 文档: {e}")

    if not doc.tables:
        raise NipptDocxParseError(f"{filename}: 文档中没有表格")

    # --- table0: seq + cilent_code ---
    table0 = []
    for row in doc.tables[0].rows:
        for cell in row.cells:
            table0.append(cell.text.strip())
    try:
        seq, cilent_code = extract_seq_code(table0)
    except IndexError:
        raise NipptDocxParseError(f"{filename}: 表格0结构异常")

    seq = seq.replace("Cliente nº ", "").replace(" ", "").replace("_", "").replace("Clientenº:", "")
    if not seq:
        raise NipptDocxParseError(f"{filename}: 无法提取 Seq")

    # Sales/Agent 从 cilent_code
    m = re.search(r"^VGBR([A-Z]+)\d+$", cilent_code or "")
    source = m.group(1) if m else ""

    # --- table13: 价格/日期/地点表 ---
    silces = []
    for i in range(len(doc.tables)):
        try:
            first = doc.tables[i].rows[0].cells[0].text
        except IndexError:
            continue
        if first == "Preço do teste:":
            silces.append(i)
        if first == "Local da coleta:":
            silces.append(i)
    list_3 = [doc.tables[i] for i in silces]
    table13 = []
    for table in list_3:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                table13.append(t if t else "")

    def t13(idx):
        return table13[idx] if len(table13) > idx else ""

    # 报告日期
    report_due_date = "".join(re.findall(r"resultados(.+?)$", t13(15)))
    report_due_date = report_due_date.replace("(DD/MM/YY)", "").replace(":", "").strip()
    client_name = title_pt(t13(9))
    wd = extract_days(t13(15))
    collect_locate = replace_collect_locate(t13(-7) if len(table13) >= 7 else "")
    gestational_week = extract_gestational_week(table13)
    observacoes = table13[-1] if table13 else ""

    # --- 人员表 ---
    people_table = None
    for i in range(len(doc.tables)):
        try:
            if doc.tables[i].rows[0].cells[0].text == "Nome Gestante":
                people_table = doc.tables[i]
        except IndexError:
            continue
    if people_table is None:
        raise NipptDocxParseError(f"{filename}: 找不到人员表(Nome Gestante)")
    mother, fathers = parse_people_table(people_table)

    # --- 孕妇信息 ---
    fe_name = title_pt(safe_get(mother, 0))
    fe_name = fe_name.replace("+", " / ").replace(" fta", " FTA")
    rg_f = safe_get(mother, 1)
    date_of_birth = safe_get(mother, 2)
    if re.match(r"\d{1,2}/\d{1,2}/\d{4}", date_of_birth or ""):
        fe_collect_date = safe_get(mother, 4)
    else:
        fe_collect_date = safe_get(mother, 3)
    fe_collect_date = fe_collect_date.replace("/24", "/2024")

    # --- 性别鉴定 ---
    gender_info = extract_gender_info(doc)

    # --- NIPT 判定 ---
    test_item = get_test_item(observacoes, seq)
    if test_item in ("Basic", "Plus"):
        return {"_nipt": True, "test_item": test_item}

    # --- 无父样本时至少输出孕妇行 ---
    if not fathers:
        fathers = [[""] * 6]

    # --- 每名疑父一条记录，与 0813 脚本一致；此处合并为 fathers 数组 ---
    father_list = []
    for f_idx, frow in enumerate(fathers):
        is_first_father = (f_idx == 0)
        row_seq = seq if is_first_father else seq + "HB"

        father_name = title_pt(safe_get(frow, 0)).replace("+", " / ")
        if not father_name or father_name == "-":
            continue  # 无父样本行跳过疑父

        father_list.append({
            "name": father_name,
            "id_card": safe_get(frow, 1),
            "sample_types": map_sample_types(safe_get(frow, 3)),
        })

    # price/balance 仅首名疑父行有效（脚本逻辑）
    price = t13(1).replace("R$", "").strip()
    balance = t13(5).replace("R$", "").strip()

    return {
        "seq": seq,
        "mother_name": fe_name or client_name,
        "mother_dob": date_of_birth,
        "mother_id_card": rg_f,
        "gestational_age_weeks": gestational_week,
        "collection_date": fe_collect_date,
        "fathers": father_list,
        "sales_person": source,
        "price": price,
        "balance": balance,
        "gender_info": gender_info,
        "notes": observacoes,
        "expected_completion": report_due_date,
    }
